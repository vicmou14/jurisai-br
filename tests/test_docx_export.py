from io import BytesIO

from docx import Document

from app.services.docx_export import export_docx


def test_export_docx_returns_valid_document():
    filename, data = export_docx("Contestação", "Texto da peça jurídica.", "office")
    assert filename.endswith(".docx")
    assert len(data) > 100
    document = Document(BytesIO(data))
    assert any("CONTESTAÇÃO" in paragraph.text for paragraph in document.paragraphs)
    assert any("Texto da peça jurídica." in paragraph.text for paragraph in document.paragraphs)
