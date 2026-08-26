from __future__ import annotations

from datetime import datetime
from io import BytesIO
import re
from uuid import uuid4

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


def _safe_filename(value: str) -> str:
    cleaned = re.sub(r"[^a-zA-Z0-9._-]+", "_", value).strip("_")
    return cleaned or f"peca_{uuid4().hex[:8]}"


def _configure(document: Document, profile: str) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.space_after = Pt(0)
    if profile == "coder":
        section.header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_title(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(12)


def export_docx(title: str, content: str, profile: str = "office") -> tuple[str, bytes]:
    document = Document()
    _configure(document, profile)
    _add_title(document, title)

    for block in [item.strip() for item in content.split("\n\n") if item.strip()]:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Cm(1.25)
        for index, line in enumerate(block.splitlines()):
            run = paragraph.add_run(line)
            if index < len(block.splitlines()) - 1:
                run.add_break()

    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(f"JurisAI-BR — documento gerado em {datetime.now().strftime('%d/%m/%Y')}")

    output = BytesIO()
    document.save(output)
    filename = _safe_filename(title) + ".docx"
    return filename, output.getvalue()
